from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import time
import os

import urllib.request
import pymysql
from datetime import date, timedelta
from docx import Document

options = Options()
options.experimental_options["prefs"] = {'profile.default_content_settings' : {"images": 2}, 
                                        'profile.managed_default_content_settings' :  {"images": 2}}
driver = webdriver.Chrome(os.getcwd()+"/chromedriver", chrome_options=options)

connection = pymysql.connect (host='localhost',
                              user='root',
                              password='password',
                              db='new_media',
                              cursorclass=pymysql.cursors.DictCursor)
                              
yesterday = str(date.today() - timedelta(days = 1))

document = Document()
document.add_heading(yesterday + '新媒體報導整理', 0)
table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '標題'
hdr_cells[1].text = '分享數'
hdr_cells[2].text = 'Hashtag'
hdr_cells[3].text = '網站'

try:
    with connection.cursor() as cursor:
        driver.get('https://buzzorange.com/techorange/')
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)
        sourceCode = BeautifulSoup(driver.page_source, "html.parser")
        articles = sourceCode.find_all('article')
        for article in articles:
            date = article.select('time.published')[0].text.replace('/', '-')
            title = article.select('h4.entry-title')[0].text
            share = article.select('span.shareCount')[0].text
            if share.find('K') != -1:
                share = float(share.split(' ')[0]) * 1000
            else:
                share = share.split(' ')[0]
            if date == yesterday:     
                print(date)
                print(title)
                print(share)
                
                cursor = connection.cursor()
                sql = ''' INSERT INTO new_media.news(title, date, share, brand) VALUES('{}', '{}', '{}', '{}')
                      '''.format(title, date, share, '科技報橘')
                cursor.execute(sql)
                connection.commit()
                
                row_cells = table.add_row().cells
                row_cells[0].text = title
                row_cells[1].text = share
                row_cells[2].text = ''
                row_cells[3].text = '科技報橘'


        for i in range(1, 3):
            driver.get('https://www.inside.com.tw/?page=' + str(i))
            sourceCode = BeautifulSoup(driver.page_source, "html.parser")
            article_list = sourceCode.find_all('div','post_list-list_style')[0]
            articles = article_list.find_all('div','post_list_item')
            for article in articles:
                date = article.select('li.post_date')[0].text.strip().replace('/', '-')
                title = article.select('a.js-auto_break_title')[0].text
                selected_tags = article.select('a.hero_slide_tag')
                tags = ''
                for tag in selected_tags:
                    tags += tag.text + ','
                if date == yesterday:
                    print(tags)
                    print(date)
                    print(title) 

                    cursor = connection.cursor()
                    sql = ''' INSERT INTO new_media.news(title, date, tag, brand) VALUES('{}', '{}', '{}', '{}')
                          '''.format(title, date, tags, 'inside')
                    cursor.execute(sql)
                    connection.commit()
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = title
                    row_cells[1].text = ''
                    row_cells[2].text = tags
                    row_cells[3].text = 'inside'


            driver.get('https://technews.tw/page/'+str(i)+'/')
            sourceCode = BeautifulSoup(driver.page_source, "html.parser")
            articles = sourceCode.find_all('header','entry-header')
            for article in articles:
                title = article.select('h1.entry-title')[0].text
                date = article.select('span.body')[1].text
                iframe = article.select('iframe')[1]
                response = urllib.request.urlopen(iframe.attrs['src'])
                iframe_soup = BeautifulSoup(response)
                share = iframe_soup.select('span#u_0_2')[0].text
                selected_tags = article.select('a[href*="category"]')
                tags = ''
                for tag in selected_tags:
                    tags += tag.text + ','
                date = date.replace(' 年 ', '-');
                date = date.replace(' 月 ', '-');
                date = date.replace(' 日 ', ' ');
                if date[0:10] == yesterday:
                    print(date)
                    print(title)
                    print(share)
                    print(tags)
                    sql = ''' INSERT INTO new_media.news(title, date, share, tag, brand) 
                              VALUES('{}', '{}', '{}', '{}', '{}')
                          '''.format(title, date, share, tags, '科技新報')
                    cursor = connection.cursor()
                    cursor.execute(sql)
                    connection.commit()
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = title
                    row_cells[1].text = share
                    row_cells[2].text = tags
                    row_cells[3].text = '科技新報'
        
    document.save(yesterday + '新媒體報告.docx')
    driver.close()
    connection.close()

except Exception as e:
    print(e)
    driver.close()
    connection.close()